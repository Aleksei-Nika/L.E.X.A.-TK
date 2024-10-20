import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from docx.oxml.shared import qn  # feel free to move these out
from docx.oxml import OxmlElement

alignment_paragraph = {'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
                       'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
                       'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
                       'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
                       'DISTRIBUTE': WD_ALIGN_PARAGRAPH.DISTRIBUTE}


def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def create_doc():
    doc = docx.Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Mm(20.4)
    section.right_margin = Mm(10)
    return doc


def save_doc(doc):
    doc.save('example.doc')


def paragraph_point(doc, para_alignment):
    para_point = doc.add_paragraph()
    para_point_format = para_point.paragraph_format
    para_point_format.space_before = Pt(0)
    para_point_format.space_after = Pt(0)
    para_point_format.line_spacing = Pt(0)
    para_point_format.alignment = alignment_paragraph[para_alignment]
    return para_point


def run_point(para_point, text, font_name, font_size, bold_font, italic_font, underline_font):
    run = para_point.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold_font
    run.italic = italic_font
    run.underline = underline_font


def table_content_point(doc, para_alignment, text, font_name, font_size, bold_font, italic_font, underline_font):
    table = doc.add_table(rows=len(text), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for row in range(len(text)):
        cell = table.cell(row, 0)
        cell.width = Cm(18)
        cell.height = Pt(0)
        para_cell = cell.paragraphs[0]
        para_cell_format = para_cell.paragraph_format()
        para_cell_format.line_spacing = Pt(0)
        para_cell_format.alignment = alignment_paragraph[para_alignment]
        para_cell_format.space_before = Pt(0)
        para_cell_format.space_after = Pt(0)
        para_cell_format.left_indent = Pt(-5)
        para_cell_format.right_indent = Pt(-5)
        run_point(para_cell, text[row], font_name, font_size, bold_font, italic_font, underline_font)
        set_cell_border(cell, top={"sz": 0, "val": "None", "color": "#000000"},
                        start={"sz": 0, "val": "None", "color": "#000000"},
                        end={"sz": 0, "val": "None", "color": "#000000"}, )



def table_content_and_explanation_point(doc, content_alignment, content_text, content_font_name, content_font_size,
                                        content_bold_font, content_italic_font, content_underline_font, expl_alignment,
                                        expl_text, expl_font_name, expl_font_size, expl_bold_font, expl_italic_font,
                                        expl_underline_font):
    pass

#text='11111111111111111111111111111111111111111111111111111111111111111111111111111111111'

#print(f'{(len(text)*0.0179751369574378*12) > 17.96}, {len(text)*0.0179751369574378*12} - {17.96}')



doc = create_doc()
text = 'Объект капитального строительства:'
para = paragraph_point(doc, 'LEFT')
run_point(para, text, 'Times New Roman', 11, True, False, False)

name_object = '''"Многоэтажная жилая застройка МО Станционный сельсовет, Новосибирского района, Новосибирской области. Жилой район "Приозерный". Квартал №2. Многоквартирный многоэтажный жилой дом №15 с помещениями общественного назначения - 2 этап, по адресу: РФ, область Новосибирская, район Новосибирский, Станционный сельсовет, п.Садовый, мкрн.Приозерный"'''
para = paragraph_point(doc, 'DISTRIBUTE')
run_point(para, name_object, 'Times New Roman', 11, False, False, True)

text = '(наименование проектной документации, почтовый или строительный адрес объекта капитального строительства)'
para = paragraph_point(doc, 'CENTER')
run_point(para, text, 'Times New Roman', 6, False, True, False)

text = 'Застройщик технический заказчик, лицо ответственное за эксплуатацию здания, сооружения или региональный оператор:'
para = paragraph_point(doc, 'LEFT')
run_point(para, text, 'Times New Roman', 11, True, False, False)

developer = '''Общество с ограниченной ответственностью Специальзированный Застройщик "Энергострой", ОГРН 1185476100039, ИНН 5410077581, 630061, г.Новосибирск, ул.Тюленина, 26, офис 215, т.347-81-00, Ассоциация Региональное Отраслевое Объединение Работодателей "Саморегулируемая организация Строителей Сибирского региона", ОГРН 1095400000189, ИНН 54065222247'''
para = paragraph_point(doc, 'DISTRIBUTE')
run_point(para, developer, 'Times New Roman', 11, False, False, True)

text = '(фамилия, имя, отчество (последнее-при наличии), адресс места жительства, ОРГНИП, ИНН индивидуального предпринимателя, полное или сокращенное наименование, ОГРН, ИНН, адрес юридического лица в пределах его местонахождения, телефон/факс, полное и (или) сокращенное наименование, ОГРН, ИНН саморегулируемой организации, членом которой является указанное юридическое лицо или индивидуальный предпрениматель, (за исключением случаев, когда членство в саморегулируемых организациях в области строительства, реконструкции, капитального ремонта  объектов капитального строительства не требуется, фамилия, имя отчество (последнее-при наличии), паспортные данные, адресс места жительства телефонн/факс - для физических лиц, не являющихся индивидуальными предпренимателями)'
para = paragraph_point(doc, 'CENTER')
run_point(para, text, 'Times New Roman', 6, False, True, False)

doc.save('example1.doc')


class Point:
    def __init__(self):
        self.font_title = None
        self.size_title = None
        self.bold_title = None
        self.italic_title = None
        self.underline_title = None
        self.paragraph_title = None
        self.alignment_title = None
        self.text_title = None

        self.font_content = None
        self.size_content = None
        self.bold_content = None
        self.italic_content = None
        self.underline_content = None
        self.paragraph_content = None
        self.alignment_content = None
        self.text_list = None
        self.symbol_before_element = None
        self.symbol_split_list = None
        self.symbol_after_list = None
        self.number_element = None
        self.symbol_before_num = None
        self.symbol_after_num = None

        self.font_explanation = None
        self.size_explanation = None
        self.bold_explanation = None
        self.italic_explanation = None
        self.underline_explanation = None
        self.paragraph_explanation = None
        self.alignment_explanation = None
        self.text_explanation = None

        self.alternation_content_and_explanations = None
        self.table_format_var = None
        self.table_line_var = None
        self.table_line_heading_var = None
        self.table_line_explanation_var = None

    def set_title(self, font, size, bold, italic, underline, paragraph, alignment, text):
        self.font_title = font
        self.size_title = size
        self.bold_title = bold
        self.italic_title = italic
        self.underline_title = underline
        self.paragraph_title = paragraph
        self.alignment_title = alignment
        self.text_title = text

    def set_content(self, font, size, bold, italic, underline, paragraph, alignment, text_list,
                    symbol_before_element, symbol_split_list, symbol_after_list, number_element, symbol_before_num,
                    symbol_after_num):
        self.font_content = font
        self.size_content = size
        self.bold_content = bold
        self.italic_content = italic
        self.underline_content = underline
        self.paragraph_content = paragraph
        self.alignment_content = alignment
        self.text_list = text_list
        self.symbol_before_element = symbol_before_element
        self.symbol_split_list = symbol_split_list
        self.symbol_after_list = symbol_after_list
        self.number_element = number_element
        self.symbol_before_num = symbol_before_num
        self.symbol_after_num = symbol_after_num

    def set_explanation(self, font, size, bold, italic, underline, paragraph, alignment, text):
        self.font_explanation = font
        self.size_explanation = size
        self.bold_explanation = bold
        self.italic_explanation = italic
        self.underline_explanation = underline
        self.paragraph_explanation = paragraph
        self.alignment_explanation = alignment
        self.text_explanation = text

    def set_parameters(self, alternation, table_format_var, table_line_var, table_line_heading_var, table_line_explanation_var):
        self.alternation_content_and_explanations = alternation
        self.table_format_var = table_format_var
        self.table_line_var = table_line_var
        self.table_line_heading_var = table_line_heading_var
        self.table_line_explanation_var = table_line_explanation_var

    def get_title(self):
        return (self.font_title, self.size_title, self.bold_title, self.italic_title, self.underline_title,
                self.paragraph_title, self.alignment_title, self.text_title)

    def get_content(self):
        return (self.font_content, self.size_content, self.bold_content, self.italic_content,
                self.underline_content, self.paragraph_content, self.alignment_content, self.text_list,
                self.symbol_before_element, self.symbol_split_list, self.symbol_after_list, self.number_element,
                self.symbol_before_num, self.symbol_after_num)

    def get_explanation(self):
        return (self.font_explanation, self.size_explanation, self.bold_explanation, self.italic_explanation,
                self.underline_explanation, self.paragraph_explanation, self.alignment_explanation,
                self.text_explanation)

    def get_parameters(self):
        return (self.alternation_content_and_explanations, self.table_format_var, self.table_line_var,
                self.table_line_heading_var, self.table_line_explanation_var)

    # ВСТАВИТЬ ФУНКЦИЮ В ИНТЕРФЕЙС!!!!!
    def get_text_content(self, content, text_list, symbol_before_element, symbol_split_list, symbol_after_list,
                         number_element, symbol_before_num, symbol_after_num):

        text = ''
        if text_list == 'ROW':
            separator = '\n'
        elif text_list == 'COLUMN':
            separator = ' '

        before_element = []
        if number_element:
            for number in range(len(content)):
                before_element.append(f'{symbol_before_num}{number+1}{symbol_after_num}')
        else:
            before_element = [symbol_before_element] * (len(content)+1)

        for index in range(len(content)):
            if index != len(content):
                text += f'{before_element}{content[index+1]}{symbol_split_list}{separator}'
            else:
                text += f'{before_element}{content[index+1]}{symbol_after_list}'

        return text

class Point_Text:
    def __init__(self):
        self.font_title = None
        self.size_title = None
        self.bold_title = None
        self.italic_title = None
        self.underline_title = None
        self.paragraph_title = None
        self.alignment_title = None
        self.text_title = None

    def set_title(self, font, size, bold, italic, underline, paragraph, alignment, text):
        self.font_title = font
        self.size_title = size
        self.bold_title = bold
        self.italic_title = italic
        self.underline_title = underline
        self.paragraph_title = paragraph
        self.alignment_title = alignment
        self.text_title = text

    def get_title(self):
        return (self.font_title, self.size_title, self.bold_title, self.italic_title, self.underline_title,
                self.paragraph_title, self.alignment_title, self.text_title)


