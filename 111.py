import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from docx.oxml.shared import qn  # feel free to move these out
from docx.oxml import OxmlElement

name_object = '''"Многоэтажная жилая застройка МО Станционный сельсовет, Новосибирского района, Новосибирской области. Жилой район "Приозерный". Квартал №2. Многоквартирный многоэтажный жилой дом №15 с помещениями общественного назначения - 2 этап, по адресу: РФ, область Новосибирская, район Новосибирский, Станционный сельсовет, п.Садовый, мкрн.Приозерный"'''
developer = '''Общество с ограниченной ответственностью Специальзированный Застройщик "Энергострой", ОГРН 1185476100039, ИНН 5410077581, 630061, г.Новосибирск, ул.Тюленина, 26, офис 215, т.347-81-00, Ассоциация Региональное Отраслевое Объединение Работодателей "Саморегулируемая организация Строителей Сибирского региона", ОГРН 1095400000189, ИНН 54065222247'''

doc = docx.Document()

section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.left_margin = Mm(20.4)
section.right_margin = Mm(10)

heading_name_object1 = doc.add_paragraph()
heading_name_object = heading_name_object1.add_run('Объект капитального строительства:')
heading_name_object.font.size = Pt(11)
heading_name_object.font.name = 'Times New Roman'
heading_name_object.bold = True
p_fmt = heading_name_object1.paragraph_format
p_fmt.space_before = Pt(0)
p_fmt.space_after = Pt(0)
p_fmt.line_spacing = Pt(0)

table_name_object = doc.add_table(rows=1, cols=1)
table_name_object.alignment = WD_TABLE_ALIGNMENT.CENTER
table_name_object.style = 'Table Grid'
cell_name_object = table_name_object.cell(0, 0)
cell_name_object.width = Cm(18)
cell_name_object.height = Pt(0)
cell_name_object_paragraph = cell_name_object.paragraphs[0]
cell_name_object_run = cell_name_object_paragraph.add_run(name_object)
cell_name_object_run.underline = True
cell_name_object_run.font.name = 'Times New Roman'
cell_name_object_run.font.size = Pt(11)
p_fmt3 = cell_name_object_paragraph.paragraph_format
p_fmt3.line_spacing = Pt(0)
p_fmt3.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
p_fmt3.space_before = Pt(0)
p_fmt3.space_after = Pt(0)
p_fmt3.left_indent = Pt(-5)
p_fmt3.right_indent = Pt(-5)


def _set_cell_background(cell, fill, color=None, val=None):
    """
    @fill: Specifies the color to be used for the background
    @color: Specifies the color to be used for any foreground
    pattern specified with the val attribute
    @val: Specifies the pattern to be used to lay the pattern
    color over the background color.
    """

    cell_properties = cell._element.tcPr
    try:
        cell_shading = cell_properties.xpath('w:shd')[0]  # in case there's already shading
    except IndexError:
        cell_shading = OxmlElement('w:shd')  # add new w:shd element to it
    if fill:
        cell_shading.set(qn('w:fill'), fill)  # set fill property, respecting namespace
    if color:
        pass  # TODO
    if val:
        pass  # TODO
    cell_properties.append(cell_shading)  # finally extend cell props with shading element


# _set_cell_background(cell, '000000')

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


set_cell_border(cell_name_object,
                top={"sz": 0, "val": "None", "color": "#000000"},
                start={"sz": 0, "val": "None", "color": "#000000"},
                end={"sz": 0, "val": "None", "color": "#000000"},
                )

explanation_name_object1 = doc.add_paragraph()
explanation_name_object = explanation_name_object1.add_run(
    '(наименование проектной документации, почтовый или строительный адрес объекта капитального строительства)')
explanation_name_object.font.size = Pt(6)
explanation_name_object.font.name = 'Times New Roman'
explanation_name_object.italic = True
p_fmt2 = explanation_name_object1.paragraph_format
p_fmt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fmt2.space_before = Pt(0)
p_fmt2.space_after = Pt(0)
p_fmt2.line_spacing = Pt(0)

heading_developer1 = doc.add_paragraph()
heading_developer = heading_developer1.add_run(
    'Застройщик технический заказчик, лицо ответственное за эксплуатацию здания, сооружения или региональный оператор:')
heading_developer.bold = True
heading_developer.font.size = Pt(11)
heading_developer.font.name = 'Times New Roman'
p_fmt_developer1 = heading_developer1.paragraph_format
p_fmt_developer1.space_before = Pt(0)
p_fmt_developer1.space_after = Pt(0)
p_fmt_developer1.line_spacing = Pt(0)

table_developer = doc.add_table(rows=1, cols=1)
table_developer.alignment = WD_TABLE_ALIGNMENT.CENTER
table_developer.style = 'Table Grid'
cell_developer = table_developer.cell(0, 0)
cell_developer.width = Cm(18)
cell_developer.height = Pt(0)
cell_developer_paragraph = cell_developer.paragraphs[0]
cell_run = cell_developer_paragraph.add_run(developer)
cell_run.underline = True
cell_run.font.size = Pt(11)
cell_run.font.name = 'Times New Roman'
p_fmt_developer = cell_developer_paragraph.paragraph_format
p_fmt_developer.line_spacing = Pt(0)
p_fmt_developer.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
p_fmt_developer.space_before = Pt(0)
p_fmt_developer.space_after = Pt(0)
p_fmt_developer.left_indent = Pt(-5)
p_fmt_developer.right_indent = Pt(-5)
set_cell_border(cell_developer,
                top={"sz": 0, "val": "None", "color": "#000000"},
                start={"sz": 0, "val": "None", "color": "#000000"},
                end={"sz": 0, "val": "None", "color": "#000000"},
                )

explanation_developer1 = doc.add_paragraph()
explanation_developer = explanation_developer1.add_run(
    '(фамилия, имя, отчество (последнее-при наличии), адресс места жительства, ОРГНИП, ИНН индивидуального предпринимателя, полное или сокращенное наименование, ОГРН, ИНН, адрес юридического лица в пределах его местонахождения, телефон/факс, полное и (или) сокращенное наименование, ОГРН, ИНН саморегулируемой организации, членом которой является указанное юридическое лицо или индивидуальный предпрениматель, (за исключением случаев, когда членство в саморегулируемых организациях в области строительства, реконструкции, капитального ремонта  объектов капитального строительства не требуется, фамилия, имя отчество (последнее-при наличии), паспортные данные, адресс места жительства телефонн/факс - для физических лиц, не являющихся индивидуальными предпренимателями)')
explanation_developer.font.size = Pt(6)
explanation_developer.font.name = 'Times New Roman'
explanation_developer.italic = True
p_fmt_explanation_developer = explanation_developer1.paragraph_format
p_fmt_explanation_developer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fmt_explanation_developer.space_before = Pt(0)
p_fmt_explanation_developer.space_after = Pt(0)
p_fmt_explanation_developer.line_spacing = Pt(0)


class Export_Akt_Word:
    def __init__(self):
        self.doc = docx.Document()

        self.section = doc.sections[0]
        self.section.page_height = Cm(29.7)
        self.section.page_width = Cm(21.0)
        self.section.left_margin = Mm(20.4)
        self.section.right_margin = Mm(10)

        class Point:
            def __init__(self):
                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run('Объект капитального строительства:')
                heading_run.font.size = Pt(11)
                heading_run.font.name = 'Times New Roman'
                heading_run.bold = True
                heading_run.italic = True
                heading_run.underline = True
                p_fmt = heading_para.paragraph_format
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(0)
                p_fmt.line_spacing = Pt(0)

                table_name_object = doc.add_table(rows=1, cols=1)
                table_name_object.alignment = WD_TABLE_ALIGNMENT.CENTER
                table_name_object.style = 'Table Grid'
                cell_name_object = table_name_object.cell(0, 0)
                cell_name_object.width = Cm(18)
                cell_name_object.height = Pt(0)
                cell_name_object_paragraph = cell_name_object.paragraphs[0]
                cell_name_object_run = cell_name_object_paragraph.add_run(name_object)
                cell_name_object_run.underline = True
                cell_name_object_run.font.name = 'Times New Roman'
                cell_name_object_run.font.size = Pt(11)
                p_fmt3 = cell_name_object_paragraph.paragraph_format
                p_fmt3.line_spacing = Pt(0)
                p_fmt3.alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE
                p_fmt3.space_before = Pt(0)
                p_fmt3.space_after = Pt(0)
                p_fmt3.left_indent = Pt(-5)
                p_fmt3.right_indent = Pt(-5)


table_new = doc.add_table(rows=2, cols=2)
table_new.style = 'Table Grid'
table_new.cell(0, 0).width = Pt(5)
set_cell_border(table_new.cell(0, 0),
                top={"sz": 0, "val": "None", "color": "#000000"},
                start={"sz": 0, "val": "None", "color": "#000000"},
                end={"sz": 0, "val": "None", "color": "#000000"},
                bottom={"sz": 0, "val": "None", "color": "#000000"},
                )
set_cell_border(table_new.cell(0, 1),
                top={"sz": 0, "val": "None", "color": "#000000"},
                start={"sz": 0, "val": "None", "color": "#000000"},
                end={"sz": 0, "val": "None", "color": "#000000"},
                )
set_cell_border(table_new.cell(1, 0),
                top={"sz": 0, "val": "None", "color": "#000000"},
                start={"sz": 0, "val": "None", "color": "#000000"},
                end={"sz": 0, "val": "None", "color": "#000000"},
                )
set_cell_border(table_new.cell(1, 1),
                top={"sz": 0, "val": "None", "color": "#000000"},
                start={"sz": 0, "val": "None", "color": "#000000"},
                end={"sz": 0, "val": "None", "color": "#000000"},
                )
table_new.cell(1, 0).merge(table_new.cell(1, 1))

doc.save('example.doc')
